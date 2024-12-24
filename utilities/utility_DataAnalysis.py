import ast

import numpy as np
import pandas as pd
from docx import Document
from utils.ComputationalModeling import bayes_factor


def mean_AIC_BIC(df):
    print(f"AIC: {df['AIC'].mean()}")
    print(f"BIC: {df['BIC'].mean()}")


# Function to create Bayes factor matrix
def create_bayes_matrix(simulations, file_name):

    def format_large_numbers(num):
        if num > 1e100:
            return f">10^100"
        if 1e100 > num > 1000:
            exponent = np.floor(np.log10(num))
            mantissa = num / 10 ** exponent
            return f"{mantissa:.3f} * 10^{exponent:.0f}"
        if num < 0.001:
            return f"<0.001"
        else:
            return f"{num:.3f}"

    # Function to add a dataframe to the document
    def add_df_to_doc(df, title):
        doc = Document()
        doc.add_heading(title, level=1)
        table = doc.add_table(df.shape[0] + 1, df.shape[1] + 1)  # Add an extra column for the row names

        # Add the column names
        for j in range(df.shape[-1]):
            table.cell(0, j + 1).text = df.columns[j]  # Shift the column names to the right

        # Add the row names and data
        for i in range(df.shape[0]):
            table.cell(i + 1, 0).text = df.index[i]  # Add the row name
            for j in range(df.shape[-1]):
                table.cell(i + 1, j + 1).text = str(df.values[i, j])  # Shift the data to the right

        doc.add_paragraph("\n")

        # Save the document
        doc.save(f"./data/DataFitting/BayesFactor/{title}.docx")

    model_names = list(simulations.keys())
    bayes_matrix = pd.DataFrame(index=model_names, columns=model_names)
    for null_model in model_names:
        for alternative_model in model_names:
            if null_model != alternative_model:
                bayes_matrix.loc[null_model, alternative_model] = bayes_factor(simulations[null_model], simulations[alternative_model])

    bayes_matrix = bayes_matrix.T

    add_df_to_doc(bayes_matrix.map(format_large_numbers), file_name)
    return bayes_matrix.map(format_large_numbers)


def safely_evaluate(x):
    if isinstance(x, list):
        return x
    try:
        # Try to safely evaluate the string to a list using ast.literal_eval
        return ast.literal_eval(x)
    except (ValueError, SyntaxError):
        # If it's not evaluable (e.g., a number), just return as is
        return [x]  # Wrap in a list for consistency


def process_chosen_prop(results, data, sub=False, values=None):

    pd.options.mode.chained_assignment = None

    # Loop through each value in the list of values
    for value in values:
        # Apply safely_evaluate and explode the column
        results[value] = results[value].apply(safely_evaluate)
        process_chosen = results[value].explode()
        # Create a new column in data for each value processed
        data[value] = process_chosen.values

    # Check if there is a TrialType column, create if absent
    if 'TrialType' not in data.columns:
        mapping = {0: 'AB', 1: 'CD', 2: 'CA', 3: 'CB', 4: 'AD', 5: 'BD'}
        data['TrialType'] = data['SetSeen.'].map(mapping)

    # Process chosen proportions and combine into one DataFrame
    if sub:
        process_chosen_df = data.groupby(['Subnum', 'TrialType'])[values].apply(
            lambda x: x.value_counts(normalize=True)).reset_index()
    else:
        process_chosen_df = data.groupby('TrialType')[values].apply(
            lambda x: x.value_counts(normalize=True)).unstack().fillna(0).reset_index()

    return data, process_chosen_df


def count_choices(data):
    count = data.groupby(['Subnum', 'TrialType', 'KeyResponse']).size().reset_index(name='Count')
    total_count = data.groupby(['Subnum', 'TrialType']).size().reset_index(name='TotalCount')
    data = pd.merge(count, total_count, on=['Subnum', 'TrialType'])
    data['optimal_ratio'] = data['Count'] / data['TotalCount']
    data = data[data['KeyResponse'].isin([0, 2])]
    data.drop(columns=['Count', 'TotalCount', 'KeyResponse'], inplace=True)
    return data


def summary_choices(data):
    data = data.groupby(['Subnum', 'TrialType', 'KeyResponse'])['Reward'].mean().reset_index()
    # transform key response to option index
    data['KeyResponse'] = data['KeyResponse'].replace({0: 'A', 1: 'B', 2: 'C', 3: 'D'})
    data_pivot = data.pivot_table(index=['Subnum'], columns='KeyResponse', values='Reward').reset_index()
    # if any option contains NaN, fill it with 0
    data_pivot.fillna(0, inplace=True)
    data_pivot.loc[:, 'ratio_AB'] = data_pivot['A'] / (data_pivot['A'] + data_pivot['B'])
    data_pivot.loc[:, 'ratio_CD'] = data_pivot['C'] / (data_pivot['C'] + data_pivot['D'])
    data_pivot.loc[:, 'ratio_CA'] = data_pivot['C'] / (data_pivot['A'] + data_pivot['C'])
    return data_pivot


def calculate_mean_squared_error(error_list):
    squared_error = [x ** 2 for x in error_list]
    return sum(squared_error) / len(error_list)


# ======================================================================================================================
# Functions for generating fitting summary
# ======================================================================================================================
def weight_calculation(group, indices):
    for index in indices:
        index_min = group[index].min()
        delta_index = group[index] - index_min
        column_name = f"{index}_weight"
        group[column_name] = np.exp(-0.5 * delta_index) / np.sum(np.exp(-0.5 * delta_index))
    return group


def extract_all_parameters(param_str):
    """
    Extracts all numerical values from a parameter string and returns them as a list of floats.

    Parameters:
    param_str (str): A string containing numerical values.

    Returns:
    list: A list of floats extracted from the string.
    """
    if isinstance(param_str, str):
        return [float(x) for x in param_str.strip('[]').split()]
    return []


def aggregate_parameters(fitting_results):
    """
    Extracts and calculates the mean of all fitting parameters for each model.

    Parameters:
    fitting_results (dict): A dictionary containing fitting results for different models.

    Returns:
    pd.DataFrame: A DataFrame containing the mean of all fitting parameters for each model.
    """
    # Initialize a dictionary to store results
    results = {}

    for key in fitting_results:
        # Extract all parameters for each model and store them in a list of lists
        all_params = fitting_results[key]['best_parameters'].apply(extract_all_parameters).tolist()

        # Transpose the list of lists to group parameters together
        all_params_transposed = list(map(list, zip(*all_params)))

        # Calculate the mean for each parameter
        param_means = [np.nanmean(params) for params in all_params_transposed]

        # Store the means in the results dictionary
        results[key] = param_means

    # Create a DataFrame from the results dictionary
    # The number of parameters might vary, so we dynamically create column names
    max_params = max(len(params) for params in results.values())
    column_names = [f'param_{i + 1}_mean' for i in range(max_params)]
    results_df = pd.DataFrame.from_dict(results, orient='index', columns=column_names)

    return results_df


def fitting_summary_generator(results, models, indices):
    fitting_summary = {}

    for key in results:
        bic_model = results[key]['BIC'].mean()
        aic_model = results[key]['AIC'].mean()
        fitting_summary[key] = {'AIC': aic_model, 'BIC': bic_model}

    fitting_summary = pd.DataFrame.from_dict(fitting_summary, orient='index', columns=indices)
    fitting_summary.reset_index(inplace=True)

    # select the models to be compared
    pattern = '|'.join(models)
    fitting_summary = fitting_summary[fitting_summary['index'].str.contains(pattern) &
                                      ~fitting_summary['index'].str.contains('uncertainty')]

    # divide by condition
    fitting_summary['condition'] = fitting_summary['index'].str.split('_').str[-2]

    result = fitting_summary.groupby('condition', group_keys=False).apply(weight_calculation, indices)

    params_df = aggregate_parameters(results)
    params_df.reset_index(inplace=True)
    result = pd.merge(result, params_df, on='index', how='inner')
    result = result.dropna(axis=1, how='all')

    return result


# Save the DataFrame as a Word document
def save_df_to_word(df, filename):
    document = Document()
    document.add_heading('Merged DataFrame', level=1)

    # Add a table to the document
    table = document.add_table(rows=1, cols=len(df.columns) + 1)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Index'
    for i, column in enumerate(df.columns):
        hdr_cells[i + 1].text = column

    # Add the DataFrame content to the table
    for idx, row in df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        for i, value in enumerate(row):
            row_cells[i + 1].text = str(value)

    # Save the document
    document.save(f"./data/DataFitting/{filename}")


# ======================================================================================================================
# Functions for individual parameters
# ======================================================================================================================
def expand_lists(df, param_cols):
    expanded_data = []
    for idx, row in df.iterrows():
        num_values = len(row[param_cols[0]])  # Assumes all param columns have the same length
        expanded_rows = pd.DataFrame({
            'index': [row['index']] * num_values,
            **{param: row[param] for param in param_cols}
        })
        expanded_data.append(expanded_rows)
    return pd.concat(expanded_data, ignore_index=True)


def individual_param_generator(results, param_cols):
    individual_param = {}

    for key in results:
        # Extract all parameters for each model and store them in a list of lists
        all_params = results[key]['best_parameters'].apply(extract_all_parameters).tolist()

        # Transpose the list of lists to group parameters together
        all_params_transposed = list(map(list, zip(*all_params)))

        individual_param[key] = all_params_transposed

    # Create a DataFrame from the result dictionary
    # The number of parameters might vary, so we dynamically create column names
    max_params = max(len(params) for params in individual_param.values())
    column_names = [f'param_{i + 1}' for i in range(max_params)]
    individual_param_df = pd.DataFrame.from_dict(individual_param, orient='index', columns=column_names)
    individual_param_df.reset_index(inplace=True)

    # Expand the DataFrame to have one row per parameter
    expanded_df = expand_lists(individual_param_df, param_cols)
    expanded_df = expanded_df.dropna(axis=1, how='all')

    # separate the index to get condition and model
    expanded_df['condition'] = expanded_df['index'].str.split('_').str[-2]
    expanded_df['model'] = expanded_df['index'].str.split('_').str[0]
    expanded_df = expanded_df.drop(columns=['index'])

    return expanded_df


# Function to calculate the AIC or BIC difference from the reference model within each group
def calculate_difference(group, ref_model):
    ref_aic = group[group['model'] == ref_model]['AIC']
    ref_bic = group[group['model'] == ref_model]['BIC']
    if not ref_aic.empty and not ref_bic.empty:
        ref_aic = ref_aic.iloc[0]
        ref_bic = ref_bic.iloc[0]
        group['AIC_diff'] = group['AIC'] - ref_aic
        group['BIC_diff'] = group['BIC'] - ref_bic
    else:
        group['AIC_diff'] = None
        group['BIC_diff'] = None
    return group


# ======================================================================================================================
# Functions for individual parameters
# ======================================================================================================================
def option_mean_calculation(data):
    training_data = data[data['Trial'] <= 150]
    mean = training_data.groupby('choice')['points'].mean()

    # keep track of the dynamic mean for each participant
    training_data.loc[:, 'cumulative_mean'] = (training_data.groupby('subnum')['points']
                                               .expanding().mean().reset_index(level=0, drop=True))
    training_data.loc[:, 'above_average'] = training_data['points'] > training_data['cumulative_mean']

    # Calculate the percentage of above-average outcomes for each choice
    above_average_counts = training_data.groupby(['choice', 'above_average']).size().unstack(fill_value=0)
    above_average_percentage = above_average_counts.div(above_average_counts.sum(axis=1), axis=0)

    result_dict = {'mean': mean, 'above_average_counts': above_average_counts,
                   'above_average_percentage': above_average_percentage}

    return result_dict








